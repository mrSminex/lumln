"""
Извлечение текста из файлов, присланных админом в Telegram (состав формулы).

Поддерживаются: .txt, .docx, .pdf.
Если библиотека для формата не установлена или текст не удалось извлечь,
вызывающий код должен попросить администратора прислать текст вручную.
"""

import io

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 МБ — более чем достаточно для текстового состава
SUPPORTED_EXTENSIONS = (".txt", ".docx", ".pdf")


class FileExtractionError(Exception):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Извлекает текст из файла по его расширению.
    Бросает FileExtractionError с понятным для пользователя сообщением при проблемах.
    """
    name = (filename or "").lower()

    if name.endswith(".txt"):
        return _extract_txt(file_bytes)
    if name.endswith(".docx"):
        return _extract_docx(file_bytes)
    if name.endswith(".pdf"):
        return _extract_pdf(file_bytes)

    raise FileExtractionError(
        "Неподдерживаемый формат файла. Пришлите состав в .txt, .docx, .pdf или текстом."
    )


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "utf-16"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise FileExtractionError("Не удалось прочитать текстовый файл (неизвестная кодировка).")

    text = text.strip()
    if not text:
        raise FileExtractionError("Файл пустой.")
    return text


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise FileExtractionError(
            "Обработка .docx недоступна на сервере. Пришлите состав в .txt или текстом."
        )
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception:
        raise FileExtractionError("Не удалось открыть .docx файл — возможно, он повреждён.")

    lines = [p.text for p in document.paragraphs if p.text.strip()]
    # Таблицы тоже могут содержать состав
    for table in document.tables:
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells_text:
                lines.append(" — ".join(cells_text))

    text = "\n".join(lines).strip()
    if not text:
        raise FileExtractionError("В .docx файле не найдено текста.")
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise FileExtractionError(
            "Обработка .pdf недоступна на сервере. Пришлите состав в .txt или текстом."
        )
    try:
        lines = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    lines.append(page_text.strip())
        text = "\n".join(lines).strip()
    except Exception:
        raise FileExtractionError("Не удалось открыть .pdf файл — возможно, он повреждён или это сканированное изображение без текста.")

    if not text:
        raise FileExtractionError(
            "В .pdf файле не найдено текста (возможно, это отсканированное изображение)."
        )
    return text